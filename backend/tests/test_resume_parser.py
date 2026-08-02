"""Resume text extraction must fill the search form without inventing titles."""

import io
from datetime import date

import pytest

from app.services.resume_parser import extract_profile, parse_resume


SAMPLE = """
Jane Doe
Bengaluru, India | jane@example.com

Senior Backend Engineer
Seeking: Platform Engineer

SUMMARY
Software engineer with 6 years of experience building APIs.

SKILLS
Python, FastAPI, PostgreSQL, Redis, AWS, Docker, Kubernetes, Kafka

EXPERIENCE
Acme Corp — Backend Engineer
2020 – 2024
Built payment APIs in Python and FastAPI.
"""


def labels(profile):
    return {location["label"] for location in profile.locations}


class TestExtractProfile:
    def test_reads_role_skills_years_and_location(self):
        profile = extract_profile(SAMPLE)

        assert profile.role in {"Senior Backend Engineer", "Platform Engineer", "Backend Engineer"}
        assert "Python" in profile.skills
        assert "FastAPI" in profile.skills
        assert profile.experience_years == 6.0
        assert profile.experience_bands == ["5-8"]
        assert labels(profile) & {"Bengaluru", "Bangalore", "India"}

    def test_prefers_seeking_line_as_primary_role(self):
        text = "Alex\nLooking for: Data Engineer\nSkills: Python, Spark\n"
        profile = extract_profile(text)
        assert profile.role == "Data Engineer"

    def test_years_from_date_span(self):
        text = "Engineer\n2018 - 2024\nPython\n"
        profile = extract_profile(text)
        assert profile.experience_years == 6.0
        assert profile.experience_bands == ["5-8"]

    def test_maps_one_year_to_junior_band(self):
        profile = extract_profile("Junior Developer\n1 year of experience\nPython, React\n")
        assert profile.experience_bands == ["0-2"]


class TestRoleAccuracy:
    def test_ignores_titles_belonging_to_colleagues(self):
        text = """Rahul Sharma
rahul@example.com

Backend Engineer

EXPERIENCE
Acme Corp - Backend Engineer
Jan 2022 - Dec 2024
Worked with a Product Manager and a Data Scientist to ship features.
Mentored engineers and reported to the Engineering Manager.
"""
        profile = extract_profile(text)
        assert profile.role == "Backend Engineer"
        assert profile.alternate_role not in {"Data Scientist", "Product Manager", "Engineering Manager"}

    def test_prefers_specific_title_over_bare_engineer(self):
        text = """Priya N
priya@example.com

Machine Learning Engineer

EXPERIENCE
Acme - Machine Learning Engineer
2021 - 2024
Engineer on the ranking team. Engineer of record for two services.
"""
        profile = extract_profile(text)
        assert profile.role == "Machine Learning Engineer"

    def test_keeps_seniority_from_the_header(self):
        text = """Sam Fox
sam@example.com
Staff Security Engineer

EXPERIENCE
Globex - Staff Security Engineer
2016 - 2024
"""
        profile = extract_profile(text)
        assert profile.role == "Staff Security Engineer"

    def test_normalises_sde_to_software_engineer(self):
        text = "Ravi K\nravi@example.com\nSDE II\n\nEXPERIENCE\nAmazon - SDE II\n2020 - 2024\n"
        profile = extract_profile(text)
        assert profile.role == "Software Engineer"

    def test_open_to_line_is_not_treated_as_a_title(self):
        text = """Dana Lee
dana@example.com
Frontend Engineer
Open to remote opportunities in fintech and healthcare

EXPERIENCE
Initech - Frontend Engineer
2019 - 2024
"""
        profile = extract_profile(text)
        assert profile.role == "Frontend Engineer"

    def test_warns_when_no_title_is_present(self):
        profile = extract_profile("Chris P\nchris@example.com\nI enjoy building things for the web.\n")
        assert profile.role == ""
        assert any("job title" in warning for warning in profile.warnings)


class TestSkillAccuracy:
    def test_finds_skills_that_end_in_symbols(self):
        text = "Lee W\nlee@example.com\nSoftware Engineer\n\nSKILLS\nC++, C#, .NET, Node.js, F#\n"
        profile = extract_profile(text)
        for skill in ("C++", "C#", ".NET", "Node.js"):
            assert skill in profile.skills, f"{skill} missing from {profile.skills}"

    def test_does_not_invent_skills_from_ordinary_words(self):
        text = """Morgan T
morgan@example.com
Product Manager

EXPERIENCE
Acme - Product Manager
2019 - 2024
Owned go-to-market strategy and R&D budget in Spring 2021.
Wrote copy to express the roadmap and rest of the plan.
"""
        profile = extract_profile(text)
        for ghost in ("Go", "R", "Spring", "Express", "REST APIs"):
            assert ghost not in profile.skills, f"{ghost} wrongly detected in {profile.skills}"

    def test_prefers_spring_boot_over_bare_spring(self):
        text = "Ann B\nann@example.com\nBackend Engineer\n\nSKILLS\nJava, Spring Boot, Spring, Kafka\n"
        profile = extract_profile(text)
        assert "Spring Boot" in profile.skills
        assert "Spring" not in profile.skills

    def test_collapses_aliases_onto_one_canonical_name(self):
        text = "Kay R\nkay@example.com\nData Engineer\n\nSKILLS\nGolang, Postgres, k8s, JS\n"
        profile = extract_profile(text)
        assert "Go" in profile.skills
        assert "PostgreSQL" in profile.skills
        assert "Kubernetes" in profile.skills
        assert "JavaScript" in profile.skills

    def test_keeps_react_alongside_next_js(self):
        text = "Ben H\nben@example.com\nFrontend Engineer\n\nSKILLS\nReact, Next.js, TypeScript\n"
        profile = extract_profile(text)
        assert {"React", "Next.js"} <= set(profile.skills)

    def test_picks_up_stack_items_outside_the_known_list(self):
        text = "Nia S\nnia@example.com\nFull Stack Engineer\n\nSKILLS\nSvelte, Supabase, Pinecone\n"
        profile = extract_profile(text)
        assert {"Svelte", "Supabase", "Pinecone"} <= set(profile.skills)

    def test_ignores_prose_in_the_skills_section(self):
        text = """Omar F
omar@example.com
Backend Engineer

SKILLS
Python, excellent communication skills, strong problem solving, Docker
"""
        profile = extract_profile(text)
        assert "Python" in profile.skills
        assert "Docker" in profile.skills
        assert not any("communication" in skill.lower() for skill in profile.skills)


class TestExperienceYears:
    def test_ignores_graduation_dates(self):
        text = """Ira M
ira@example.com
Software Engineer

EXPERIENCE
Acme - Software Engineer
Jan 2022 - Jan 2024

EDUCATION
B.Tech, Some University, 2014 - 2018
"""
        profile = extract_profile(text)
        assert profile.experience_years == 2.0
        assert profile.experience_bands == ["2-5"]

    def test_ignores_years_quoted_about_other_people(self):
        text = """Rahul Sharma
rahul@example.com
Backend Engineer

EXPERIENCE
Acme Corp - Backend Engineer
Jan 2020 - Jan 2024
Mentored engineers with 2 years of experience.
"""
        profile = extract_profile(text)
        assert profile.experience_years == 4.0

    def test_counts_an_internship_section_as_experience(self):
        text = """Karan M
karan@example.com
Software Engineer

INTERNSHIP
Zoho - Software Engineer Intern
Jun 2024 - Dec 2024

EDUCATION
B.Tech, Some Institute, 2021 - 2025
"""
        profile = extract_profile(text)
        assert profile.experience_years == 0.5
        assert profile.experience_bands == ["0-2"]

    def test_counts_overlapping_jobs_once(self):
        text = """Tess O
tess@example.com
Backend Engineer

EXPERIENCE
Acme - Backend Engineer
Jan 2020 - Dec 2022
Beta Inc - Backend Engineer (part time)
Jun 2021 - Dec 2022
"""
        profile = extract_profile(text)
        assert profile.experience_years == pytest.approx(2.9, abs=0.15)

    def test_handles_present_as_an_end_date(self):
        today = date.today()
        start_year = today.year - 3
        text = f"""Vik A
vik@example.com
Backend Engineer

EXPERIENCE
Acme - Backend Engineer
Jan {start_year} - Present
"""
        months = (today.year * 12 + today.month) - (start_year * 12 + 1)
        profile = extract_profile(text)
        assert profile.experience_years == pytest.approx(months / 12, abs=0.05)

    def test_reads_numeric_month_ranges(self):
        text = """Sara P
sara@example.com
Data Engineer

EXPERIENCE
Acme - Data Engineer
01/2019 - 01/2024
"""
        profile = extract_profile(text)
        assert profile.experience_years == 5.0

    def test_reads_a_labelled_total(self):
        text = """Ken J
ken@example.com
Software Engineer

SUMMARY
Total Experience: 9 years

EXPERIENCE
Acme - Software Engineer
2022 - 2024
"""
        profile = extract_profile(text)
        assert profile.experience_years == 9.0
        assert profile.experience_bands == ["8-12"]

    def test_dates_win_when_the_written_claim_is_stale(self):
        text = """Lily Q
lily@example.com
Backend Engineer

SUMMARY
Backend engineer with 3 years of experience.

EXPERIENCE
Acme - Backend Engineer
Jan 2012 - Jan 2024
"""
        profile = extract_profile(text)
        assert profile.experience_years == 12.0


class TestLocations:
    def test_reads_the_city_from_the_contact_line(self):
        text = "Asha R\nPune, India | asha@example.com\nBackend Engineer\n\nSKILLS\nPython\n"
        profile = extract_profile(text)
        assert "Pune" in labels(profile)

    def test_does_not_use_a_university_city_as_a_home(self):
        text = """Rahul Sharma
Delhi Technological University alumnus | rahul@example.com
Backend Engineer

SKILLS
Python

EDUCATION
B.Tech, Delhi Technological University, 2014 - 2018
"""
        profile = extract_profile(text)
        assert "Delhi" not in labels(profile)

    def test_drops_the_country_when_a_city_is_known(self):
        text = "Jane Doe\nBengaluru, India | jane@example.com\nBackend Engineer\n\nSKILLS\nPython\n"
        profile = extract_profile(text)
        assert labels(profile) == {"Bengaluru"}

    def test_prefers_the_longer_place_name(self):
        text = "Amit V\nNew Delhi | amit@example.com\nBackend Engineer\n\nSKILLS\nPython\n"
        profile = extract_profile(text)
        assert "New Delhi" in labels(profile)

    def test_reads_a_labelled_location_line(self):
        text = """Sana K
sana@example.com
Backend Engineer

SKILLS
Python

Preferred Location: Hyderabad
"""
        profile = extract_profile(text)
        assert "Hyderabad" in labels(profile)

    def test_flags_remote(self):
        text = "Nils B\nRemote | nils@example.com\nBackend Engineer\n\nSKILLS\nPython\n"
        profile = extract_profile(text)
        assert any(location["remote"] for location in profile.locations)


class TestMessyFormatting:
    def test_reads_bullet_glyphs_and_ligatures(self):
        text = "Ada L\nada@example.com\nBackend Engineer\n\nSKILLS\n• Python • \uFB02ask • Docker\n"
        profile = extract_profile(text)
        assert {"Python", "Flask", "Docker"} <= set(profile.skills)

    def test_rejoins_words_split_across_lines(self):
        text = "Ada L\nada@example.com\nBackend Engineer\n\nSKILLS\nKuber-\nnetes, Python\n"
        profile = extract_profile(text)
        assert "Kubernetes" in profile.skills

    def test_splits_two_column_layouts_on_wide_gaps(self):
        text = (
            "Ada L\nada@example.com\nBackend Engineer\n\n"
            "SKILLS                    EXPERIENCE\n"
            "Python                    Acme - Backend Engineer\n"
            "Docker                    Jan 2020 - Jan 2024\n"
        )
        profile = extract_profile(text)
        assert "Python" in profile.skills
        assert profile.experience_years == 4.0

    def test_reads_inline_section_headers(self):
        text = "Ada L\nada@example.com\nBackend Engineer\nSkills: Python, Docker, Terraform\n"
        profile = extract_profile(text)
        assert {"Python", "Docker", "Terraform"} <= set(profile.skills)


def _build_pdf(lines: list[str]) -> bytes:
    """A minimal single page PDF so the pypdf path is covered end to end."""
    text_ops = "".join(f"({line}) Tj T*\n" for line in lines)
    stream = f"BT /F1 12 Tf 50 750 Td 15 TL\n{text_ops}ET"
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n{body}\nendobj\n".encode("latin-1")

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode("latin-1")
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode("latin-1")
    return bytes(out)


def _build_docx(rows: list[tuple[str, str]], paragraphs: list[str]) -> bytes:
    from docx import Document

    document = Document()
    for line in paragraphs:
        document.add_paragraph(line)
    table = document.add_table(rows=len(rows), cols=2)
    for row, (left, right) in zip(table.rows, rows):
        row.cells[0].text = left
        row.cells[1].text = right
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestParseResumeFile:
    def test_reads_plain_text_upload(self):
        profile = parse_resume("resume.txt", SAMPLE.encode("utf-8"))
        assert profile.role
        assert profile.skills

    def test_reads_a_pdf_upload(self):
        pdf = _build_pdf(
            [
                "Jane Doe",
                "Pune, India | jane@example.com",
                "Backend Engineer",
                "SKILLS",
                "Python, FastAPI, Docker, Kubernetes",
                "EXPERIENCE",
                "Acme Corp - Backend Engineer",
                "Jan 2020 - Jan 2024",
            ]
        )
        profile = parse_resume("resume.pdf", pdf)
        assert profile.role == "Backend Engineer"
        assert "Python" in profile.skills
        assert profile.experience_years == 4.0

    def test_reads_a_docx_whose_content_lives_in_a_table(self):
        docx = _build_docx(
            rows=[
                ("Skills", "Python, FastAPI, Docker"),
                ("Experience", "Acme Corp - Backend Engineer, Jan 2020 - Jan 2024"),
            ],
            paragraphs=["Jane Doe", "Pune, India | jane@example.com", "Backend Engineer"],
        )
        profile = parse_resume("resume.docx", docx)
        assert profile.role == "Backend Engineer"
        assert "Python" in profile.skills
        assert "Pune" in labels(profile)

    def test_rejects_unsupported_type(self):
        with pytest.raises(ValueError, match="Unsupported"):
            parse_resume("photo.png", b"not a resume")

    def test_rejects_tiny_file(self):
        with pytest.raises(ValueError, match="enough text"):
            parse_resume("resume.txt", b"hi")

    def test_rejects_oversized_file(self):
        with pytest.raises(ValueError, match="5 MB"):
            parse_resume("resume.pdf", b"x" * (5 * 1024 * 1024 + 1))

    def test_explains_image_only_pdfs(self):
        with pytest.raises(ValueError, match="scanned"):
            parse_resume("resume.pdf", _build_pdf([]))
