"""Extract searchable job filters from an uploaded resume.

No LLM required. The resume is read as text (PDF / DOCX / plain), normalised,
then split into its sections (header, summary, skills, experience, education,
...). Each field is read from the section it actually belongs to, which keeps
stray mentions -- a colleague's job title in a bullet, a university's city, a
graduation year -- from overwriting the real answer.
"""

from __future__ import annotations

import io
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import date

# --------------------------------------------------------------------------
# Vocabulary
# --------------------------------------------------------------------------

# (canonical title, surface forms seen in resumes). Canonical form is what we
# put in the search box, so it should read like a job board listing.
_TITLE_DEFS: list[tuple[str, tuple[str, ...]]] = [
    ("Machine Learning Engineer", ("machine learning engineer", "ml engineer", "mle")),
    ("AI Engineer", ("ai engineer", "artificial intelligence engineer", "genai engineer", "generative ai engineer")),
    ("Data Scientist", ("data scientist", "applied scientist")),
    ("Data Engineer", ("data engineer", "big data engineer")),
    ("Data Analyst", ("data analyst",)),
    ("Analytics Engineer", ("analytics engineer",)),
    ("Business Analyst", ("business analyst",)),
    ("Research Engineer", ("research engineer",)),
    ("Research Scientist", ("research scientist",)),
    ("Full Stack Engineer", (
        "full stack engineer", "full-stack engineer", "fullstack engineer",
        "full stack developer", "full-stack developer", "fullstack developer",
        "full stack web developer",
    )),
    ("Frontend Engineer", (
        "frontend engineer", "front end engineer", "front-end engineer",
        "frontend developer", "front end developer", "front-end developer",
        "ui engineer", "ui developer",
    )),
    ("Backend Engineer", (
        "backend engineer", "back end engineer", "back-end engineer",
        "backend developer", "back end developer", "back-end developer",
        "server side engineer", "server-side engineer",
    )),
    ("Platform Engineer", ("platform engineer",)),
    ("Infrastructure Engineer", ("infrastructure engineer", "infra engineer")),
    ("DevOps Engineer", ("devops engineer", "dev ops engineer")),
    ("Cloud Engineer", ("cloud engineer",)),
    ("Site Reliability Engineer", ("site reliability engineer", "sre")),
    ("Security Engineer", (
        "security engineer", "application security engineer",
        "infosec engineer", "cyber security engineer", "cybersecurity engineer",
    )),
    ("Network Engineer", ("network engineer",)),
    ("Embedded Engineer", ("embedded engineer", "embedded software engineer", "firmware engineer")),
    ("iOS Engineer", ("ios engineer", "ios developer")),
    ("Android Engineer", ("android engineer", "android developer")),
    ("Mobile Engineer", (
        "mobile engineer", "mobile developer",
        "react native developer", "flutter developer",
    )),
    ("QA Engineer", (
        "qa engineer", "quality assurance engineer", "test engineer",
        "automation test engineer", "sdet", "software development engineer in test",
    )),
    ("Solutions Architect", ("solutions architect", "solution architect")),
    ("Software Architect", ("software architect", "technical architect")),
    ("Engineering Manager", ("engineering manager", "software engineering manager")),
    ("Product Manager", ("product manager", "associate product manager", "technical product manager")),
    ("Program Manager", ("program manager", "technical program manager")),
    ("Project Manager", ("project manager",)),
    ("Product Designer", ("product designer", "ux designer", "ui/ux designer", "user experience designer")),
    ("Database Administrator", ("database administrator", "dba")),
    ("Systems Engineer", ("systems engineer", "system engineer")),
    ("Software Engineer", (
        "software engineer", "software development engineer", "swe",
        "sde", "sde i", "sde ii", "sde iii", "sde 1", "sde 2", "sde 3",
        "sde-1", "sde-2", "sde-3",
    )),
    ("Software Developer", ("software developer", "application developer", "programmer analyst")),
    ("Web Developer", ("web developer",)),
    # Fallbacks: only used when nothing more specific is found.
    ("Engineer", ("engineer",)),
    ("Developer", ("developer",)),
]

GENERIC_TITLES = {"Engineer", "Developer"}

SENIORITY_FORMS = {
    "sr": "Senior",
    "sr.": "Senior",
    "senior": "Senior",
    "staff": "Staff",
    "principal": "Principal",
    "lead": "Lead",
    "jr": "Junior",
    "jr.": "Junior",
    "junior": "Junior",
    "associate": "Associate",
    "entry level": "Junior",
    "entry-level": "Junior",
}

# (canonical skill, surface forms). Forms listed in AMBIGUOUS_FORMS only count
# when they show up inside an explicit skills section.
_SKILL_DEFS: list[tuple[str, tuple[str, ...]]] = [
    ("Python", ("python", "python3")),
    ("JavaScript", ("javascript", "java script", "js", "es6")),
    ("TypeScript", ("typescript", "ts")),
    ("Java", ("java",)),
    ("Kotlin", ("kotlin",)),
    ("Swift", ("swift",)),
    ("Go", ("golang", "go lang", "go")),
    ("Rust", ("rust",)),
    ("C++", ("c++", "cpp", "c plus plus")),
    ("C#", ("c#", "c sharp", "csharp")),
    ("C", ("c",)),
    ("Ruby", ("ruby", "ruby on rails", "rails")),
    ("PHP", ("php", "laravel")),
    ("Scala", ("scala",)),
    ("R", ("r",)),
    ("MATLAB", ("matlab",)),
    ("SQL", ("sql",)),
    ("NoSQL", ("nosql",)),
    ("PostgreSQL", ("postgresql", "postgres", "psql")),
    ("MySQL", ("mysql", "mariadb")),
    ("MongoDB", ("mongodb", "mongo")),
    ("Redis", ("redis",)),
    ("Cassandra", ("cassandra",)),
    ("Elasticsearch", ("elasticsearch", "elastic search", "opensearch")),
    ("DynamoDB", ("dynamodb", "dynamo db")),
    ("SQLite", ("sqlite",)),
    ("React", ("react", "react.js", "reactjs")),
    ("Next.js", ("next.js", "nextjs")),
    ("Vue", ("vue", "vue.js", "vuejs")),
    ("Angular", ("angular", "angularjs")),
    ("Svelte", ("svelte", "sveltekit")),
    ("Redux", ("redux",)),
    ("HTML", ("html", "html5")),
    ("CSS", ("css", "css3", "sass", "scss")),
    ("Tailwind CSS", ("tailwind", "tailwindcss", "tailwind css")),
    ("Node.js", ("node.js", "nodejs", "node js")),
    ("Express", ("express.js", "expressjs", "express")),
    ("Django", ("django",)),
    ("Flask", ("flask",)),
    ("FastAPI", ("fastapi", "fast api")),
    ("Spring Boot", ("spring boot", "springboot")),
    ("Spring", ("spring framework", "spring")),
    (".NET", (".net", "dotnet", "asp.net")),
    ("GraphQL", ("graphql",)),
    ("REST APIs", ("rest api", "rest apis", "restful", "rest")),
    ("gRPC", ("grpc",)),
    ("Kafka", ("kafka",)),
    ("RabbitMQ", ("rabbitmq",)),
    ("Microservices", ("microservices", "microservice")),
    ("AWS", ("aws", "amazon web services")),
    ("GCP", ("gcp", "google cloud", "google cloud platform")),
    ("Azure", ("azure", "microsoft azure")),
    ("Docker", ("docker",)),
    ("Kubernetes", ("kubernetes", "k8s")),
    ("Terraform", ("terraform",)),
    ("Ansible", ("ansible",)),
    ("CI/CD", ("ci/cd", "cicd", "ci cd")),
    ("Jenkins", ("jenkins",)),
    ("GitHub Actions", ("github actions",)),
    ("Linux", ("linux", "unix")),
    ("Git", ("git",)),
    ("Bash", ("bash", "shell scripting")),
    ("PyTorch", ("pytorch", "torch")),
    ("TensorFlow", ("tensorflow", "tf")),
    ("scikit-learn", ("scikit-learn", "sklearn", "scikit learn")),
    ("NLP", ("nlp", "natural language processing")),
    ("LLMs", ("llm", "llms", "large language models")),
    ("RAG", ("rag", "retrieval augmented generation")),
    ("LangChain", ("langchain",)),
    ("OpenAI", ("openai",)),
    ("Spark", ("spark", "pyspark", "apache spark")),
    ("Airflow", ("airflow", "apache airflow")),
    ("dbt", ("dbt",)),
    ("Snowflake", ("snowflake",)),
    ("Databricks", ("databricks",)),
    ("Tableau", ("tableau",)),
    ("Power BI", ("power bi", "powerbi")),
    ("Pandas", ("pandas",)),
    ("NumPy", ("numpy",)),
]

# Words that are real skills but also ordinary English / part of other names.
# Requiring a skills section stops "Go-to-market", "Spring 2021" and "R&D"
# from turning into programming languages.
AMBIGUOUS_FORMS = {"go", "r", "c", "spring", "express", "rest", "ts", "tf", "torch", "swift"}

# If the more specific skill is present, drop the ambiguous parent term.
SKILL_SUPERSEDES = {"Spring Boot": "Spring"}

KNOWN_LOCATIONS: list[tuple[str, dict]] = [
    ("Bengaluru", {"city": "Bengaluru", "state": "Karnataka", "country": "India"}),
    ("Bangalore", {"city": "Bengaluru", "state": "Karnataka", "country": "India"}),
    ("Mumbai", {"city": "Mumbai", "state": "Maharashtra", "country": "India"}),
    ("Navi Mumbai", {"city": "Mumbai", "state": "Maharashtra", "country": "India"}),
    ("Pune", {"city": "Pune", "state": "Maharashtra", "country": "India"}),
    ("Hyderabad", {"city": "Hyderabad", "state": "Telangana", "country": "India"}),
    ("Chennai", {"city": "Chennai", "state": "Tamil Nadu", "country": "India"}),
    ("New Delhi", {"city": "Delhi", "state": "Delhi", "country": "India"}),
    ("Delhi", {"city": "Delhi", "state": "Delhi", "country": "India"}),
    ("Gurugram", {"city": "Gurugram", "state": "Haryana", "country": "India"}),
    ("Gurgaon", {"city": "Gurugram", "state": "Haryana", "country": "India"}),
    ("Noida", {"city": "Noida", "state": "Uttar Pradesh", "country": "India"}),
    ("Kolkata", {"city": "Kolkata", "state": "West Bengal", "country": "India"}),
    ("Ahmedabad", {"city": "Ahmedabad", "state": "Gujarat", "country": "India"}),
    ("Jaipur", {"city": "Jaipur", "state": "Rajasthan", "country": "India"}),
    ("Indore", {"city": "Indore", "state": "Madhya Pradesh", "country": "India"}),
    ("Chandigarh", {"city": "Chandigarh", "state": "Chandigarh", "country": "India"}),
    ("Kochi", {"city": "Kochi", "state": "Kerala", "country": "India"}),
    ("Coimbatore", {"city": "Coimbatore", "state": "Tamil Nadu", "country": "India"}),
    ("San Francisco", {"city": "San Francisco", "state": "California", "country": "United States"}),
    ("Bay Area", {"city": "San Francisco", "state": "California", "country": "United States"}),
    ("New York", {"city": "New York", "state": "New York", "country": "United States"}),
    ("Seattle", {"city": "Seattle", "state": "Washington", "country": "United States"}),
    ("Austin", {"city": "Austin", "state": "Texas", "country": "United States"}),
    ("Boston", {"city": "Boston", "state": "Massachusetts", "country": "United States"}),
    ("Chicago", {"city": "Chicago", "state": "Illinois", "country": "United States"}),
    ("Los Angeles", {"city": "Los Angeles", "state": "California", "country": "United States"}),
    ("Denver", {"city": "Denver", "state": "Colorado", "country": "United States"}),
    ("Atlanta", {"city": "Atlanta", "state": "Georgia", "country": "United States"}),
    ("London", {"city": "London", "state": "", "country": "United Kingdom"}),
    ("Dublin", {"city": "Dublin", "state": "", "country": "Ireland"}),
    ("Berlin", {"city": "Berlin", "state": "", "country": "Germany"}),
    ("Amsterdam", {"city": "Amsterdam", "state": "", "country": "Netherlands"}),
    ("Zurich", {"city": "Zurich", "state": "", "country": "Switzerland"}),
    ("Toronto", {"city": "Toronto", "state": "Ontario", "country": "Canada"}),
    ("Vancouver", {"city": "Vancouver", "state": "British Columbia", "country": "Canada"}),
    ("Sydney", {"city": "Sydney", "state": "New South Wales", "country": "Australia"}),
    ("Singapore", {"city": "Singapore", "state": "", "country": "Singapore"}),
    ("Dubai", {"city": "Dubai", "state": "", "country": "United Arab Emirates"}),
    ("India", {"city": "", "state": "", "country": "India"}),
    ("United States", {"city": "", "state": "", "country": "United States"}),
    ("USA", {"city": "", "state": "", "country": "United States"}),
    ("United Kingdom", {"city": "", "state": "", "country": "United Kingdom"}),
    ("UK", {"city": "", "state": "", "country": "United Kingdom"}),
    ("Canada", {"city": "", "state": "", "country": "Canada"}),
    ("Germany", {"city": "", "state": "", "country": "Germany"}),
    ("Remote", {"city": "", "state": "", "country": "", "remote": True}),
]

# A place name followed by one of these is an employer or a school, not a home.
INSTITUTION_SUFFIX = re.compile(
    r"^\s*(?:technolog\w*|institute|university|college|school|academy|"
    r"public\s+school|polytechnic|hospital|airport|metro|stock\s+exchange|"
    r"ltd\.?|limited|inc\.?|llc|llp|pvt\.?|private|corp\.?|corporation|"
    r"solutions|systems|softwares?|labs?|technologies|consultanc\w*|bank)\b",
    re.I,
)

EXPERIENCE_BANDS = [
    ("0-2", 0.0, 2.0),
    ("2-5", 2.0, 5.0),
    ("5-8", 5.0, 8.0),
    ("8-12", 8.0, 12.0),
    ("12+", 12.0, 50.0),
]

MAX_RESUME_BYTES = 5 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MAX_SKILLS = 15
MAX_PDF_PAGES = 10
# Below this a title is a passing mention rather than the candidate's own job.
MIN_ROLE_SCORE = 2.0

# Section headers, matched against a whole header line.
SECTION_PATTERNS: list[tuple[str, str]] = [
    ("summary", r"(?:professional\s+|career\s+|personal\s+|executive\s+)?"
                r"(?:summary|profile|objective|about(?:\s+me)?|overview|introduction)"),
    ("skills", r"(?:technical\s+|core\s+|key\s+|primary\s+|relevant\s+|it\s+)?"
               r"(?:skills?(?:\s+(?:set|summary|highlights))?|competencies|"
               r"tech(?:nical)?\s+stack|technologies|technical\s+proficienc\w*|"
               r"proficiencies|expertise|areas\s+of\s+expertise|"
               r"tools\s*(?:&|and|/)?\s*technologies?|languages\s*(?:&|and|/)\s*\w+)"),
    ("experience", r"(?:work\s+|working\s+|professional\s+|relevant\s+|industry\s+|"
                   r"employment\s+|career\s+)?"
                   r"(?:experience|employment(?:\s+history)?|work\s+history|"
                   r"career\s+history|professional\s+background|"
                   r"internships?(?:\s+experience)?)"),
    ("education", r"(?:education(?:al\s+(?:background|qualifications?))?|academics?|"
                  r"academic\s+(?:background|qualifications?)|qualifications?|"
                  r"scholastic\s+record)"),
    ("projects", r"(?:personal\s+|side\s+|academic\s+|key\s+|selected\s+|notable\s+)?"
                 r"projects?"),
    ("certifications", r"(?:certifications?|certificates?|licen[cs]es?|"
                       r"courses?(?:\s*work)?|coursework|training|"
                       r"online\s+courses?)"),
    ("other", r"(?:awards?(?:\s*(?:&|and)\s*\w+)?|achievements?|accomplishments|"
              r"honou?rs?(?:\s*(?:&|and)\s*\w+)?|publications?|patents?|activities|"
              r"extra[\s-]?curricular(?:\s+activities)?|interests|hobbies|languages|"
              r"references|volunteer(?:ing|\s+experience)?|"
              r"positions?\s+of\s+responsibility|leadership)"),
]

# Where each field is worth reading from. Higher means more trustworthy.
ROLE_SECTION_WEIGHT = {
    "_header": 12.0,
    "summary": 6.0,
    "experience": 5.0,
    "projects": 1.0,
    "skills": 0.5,
    "education": 0.4,
    "certifications": 0.3,
    "other": 0.3,
}

# Lines describing somebody else's job, not the candidate's.
THIRD_PARTY_CONTEXT = re.compile(
    r"\b(?:worked\s+with|collaborat\w+\s+with|partnered\s+with|alongside|"
    r"reported\s+to|reporting\s+to|mentored|mentoring|coached|managed\s+a\s+team|"
    r"team\s+of|cross[\s-]functional|stakeholders?|interviewed|hiring|recruit\w*|"
    r"liaised|coordinated\s+with|paired\s+with|supported\s+the)\b",
    re.I,
)

EDUCATION_LINE = re.compile(
    r"\b(?:b\.?\s?tech|m\.?\s?tech|b\.?\s?e\.?|b\.?\s?sc|m\.?\s?sc|b\.?\s?s\.?|m\.?\s?s\.?|"
    r"bachelor\w*|master\w*|ph\.?\s?d|mba|mca|bca|b\.?com|diploma|"
    r"university|college|institute|school|gpa|cgpa|percentage|"
    r"graduat\w*|degree|semester|class\s+(?:x|xii|10|12)\b|10th|12th)\b",
    re.I,
)

_MONTH_NUMS = {
    "january": 1, "jan": 1, "february": 2, "feb": 2, "march": 3, "mar": 3,
    "april": 4, "apr": 4, "may": 5, "june": 6, "jun": 6, "july": 7, "jul": 7,
    "august": 8, "aug": 8, "september": 9, "sept": 9, "sep": 9,
    "october": 10, "oct": 10, "november": 11, "nov": 11, "december": 12, "dec": 12,
}
_MONTH_ALT = "|".join(sorted(_MONTH_NUMS, key=len, reverse=True))

# "Jan 2020 - Mar 2022", "2018 - 2024", "03/2021 - Present"
DATE_RANGE_RE = re.compile(
    r"(?:(?P<m1>" + _MONTH_ALT + r")\.?\s*[,./-]?\s*|(?P<n1>0?[1-9]|1[0-2])\s*[/.]\s*)?"
    r"(?P<y1>(?:19|20)\d{2})"
    r"\s*(?:-{1,2}|to|until|through|till)\s*"
    r"(?:(?P<present>present|current(?:ly)?|now|ongoing|date|till\s+date|to\s+date)"
    r"|(?:(?:(?P<m2>" + _MONTH_ALT + r")\.?\s*[,./-]?\s*|(?P<n2>0?[1-9]|1[0-2])\s*[/.]\s*)?"
    r"(?P<y2>(?:19|20)\d{2})))",
    re.I,
)

# "6 years of experience", "5+ yrs of professional experience"
EXPLICIT_YEARS_RE = re.compile(
    r"(?<![\d.])(\d{1,2}(?:\.\d)?)\s*\+?\s*(?:years?|yrs?)\b"
    r"(?:\s+(?:of|in|as|with))?"
    r"(?:\s+(?:professional|industry|hands[\s-]?on|relevant|overall|total|"
    r"work|working|solid|proven|diverse|progressive))*"
    r"\s*(?:experience|exp\b)",
    re.I,
)
# "Experience: 6 years", "Total Experience - 5.5 yrs"
LABELLED_YEARS_RE = re.compile(
    r"(?:total|overall)?\s*experience\s*[:\-]\s*(\d{1,2}(?:\.\d)?)\s*\+?\s*(?:years?|yrs?)\b",
    re.I,
)

SEEKING_RE = re.compile(
    r"^[\s\-*>|]*(?:currently\s+)?"
    r"(?:seeking|looking\s+for|target(?:ing)?(?:\s+role)?|desired\s+(?:role|position)|"
    r"objective|open\s+to|aspiring|applying\s+for|interested\s+in)"
    r"\s*(?:a|an|the)?\s*[:\-]?\s*(.{3,60})$",
    re.I | re.M,
)

ROLE_NOUNS = {
    "engineer", "engineering", "developer", "scientist", "manager", "analyst",
    "designer", "architect", "administrator", "consultant", "specialist",
    "lead", "director", "intern", "programmer", "sre", "dba", "sde", "swe",
    "researcher",
}

LOCATION_LABEL_RE = re.compile(
    r"^[\s\-*>|]*(?:current\s+|preferred\s+|home\s+|base[d]?\s+)?"
    r"(?:location|locality|city|address|based\s+in|residing\s+in)\s*[:\-]\s*(.+)$",
    re.I | re.M,
)


def _build_title_matcher() -> tuple[re.Pattern[str], dict[str, str]]:
    canonical: dict[str, str] = {}
    for title, forms in _TITLE_DEFS:
        for form in forms:
            canonical[form.lower()] = title
    # Longest form first so "backend engineer" wins over bare "engineer" and we
    # never record both for the same span of text.
    alternation = "|".join(
        re.escape(form) for form in sorted(canonical, key=len, reverse=True)
    )
    seniority = "|".join(
        re.escape(form) for form in sorted(SENIORITY_FORMS, key=len, reverse=True)
    )
    pattern = re.compile(
        rf"(?<![A-Za-z])(?:(?P<sen>{seniority})\s+)?(?P<title>{alternation})(?![A-Za-z])",
        re.I,
    )
    return pattern, canonical


TITLE_RE, TITLE_CANONICAL = _build_title_matcher()


def _skill_pattern(form: str) -> re.Pattern[str]:
    """Whole-token match that also works for C++, C#, .NET and Node.js.

    A plain ``\\b`` fails on symbol-terminated skills: ``\\bC\\+\\+\\b`` needs a
    word character after the ``+``, so "C++" at the end of a line never matches.
    """
    return re.compile(
        r"(?<![A-Za-z0-9+#.])" + re.escape(form) + r"(?![A-Za-z0-9+#])",
        re.I,
    )


SKILL_MATCHERS: list[tuple[str, tuple[tuple[str, re.Pattern[str], bool], ...]]] = [
    (
        canonical,
        tuple(
            (form, _skill_pattern(form), form.lower() in AMBIGUOUS_FORMS)
            for form in forms
        ),
    )
    for canonical, forms in _SKILL_DEFS
]
def _skill_key(value: str) -> str:
    """Compare skills ignoring punctuation, so ".NET" and "NET" are one thing."""
    return re.sub(r"[^a-z0-9]", "", value.lower())


ALL_SKILL_KEYS = {
    _skill_key(form) for _, forms in _SKILL_DEFS for form in forms
} | {_skill_key(canonical) for canonical, _ in _SKILL_DEFS}


@dataclass
class ResumeProfile:
    role: str = ""
    alternate_role: str = ""
    skills: list[str] = field(default_factory=list)
    experience_years: float | None = None
    experience_bands: list[str] = field(default_factory=list)
    locations: list[dict] = field(default_factory=list)
    summary: str = ""
    warnings: list[str] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {
            "role": self.role,
            "alternate_role": self.alternate_role,
            "skills": self.skills,
            "skills_text": ", ".join(self.skills),
            "experience_years": self.experience_years,
            "experience_bands": self.experience_bands,
            "locations": self.locations,
            "summary": self.summary,
            "warnings": self.warnings,
        }


# --------------------------------------------------------------------------
# Entry points
# --------------------------------------------------------------------------


def parse_resume(filename: str, content: bytes) -> ResumeProfile:
    """Turn an uploaded resume file into search form fields."""
    if len(content) > MAX_RESUME_BYTES:
        raise ValueError("Resume is larger than 5 MB")

    name = (filename or "resume.txt").lower()
    ext = "." + name.rsplit(".", 1)[-1] if "." in name else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type. Upload a PDF, DOCX or plain text resume."
        )

    text = _extract_text(ext, content)
    if not text or len(text.strip()) < 40:
        if ext == ".pdf":
            raise ValueError(
                "Could not read enough text from that resume. If it is a scanned "
                "or image-only PDF, upload a DOCX or text version instead."
            )
        raise ValueError("Could not read enough text from that resume")

    return extract_profile(text)


def extract_profile(text: str) -> ResumeProfile:
    """Pure text -> profile. Exposed for tests without file I/O."""
    text = _normalize_text(text)
    sections = _split_sections(text)

    profile = ResumeProfile()
    roles = _extract_roles(text, sections)
    if roles:
        profile.role = roles[0]
        if len(roles) > 1:
            profile.alternate_role = roles[1]
    else:
        profile.warnings.append(
            "No job title found - enter one manually before searching."
        )

    profile.skills = _extract_skills(text, sections)
    if not profile.skills:
        profile.warnings.append("No skills recognised - add a few before searching.")

    years = _extract_years(text, sections)
    profile.experience_years = years
    if years is not None:
        profile.experience_bands = _years_to_bands(years)

    profile.locations = _extract_locations(text, sections)
    profile.summary = _short_summary(text, profile)
    return profile


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------


def _extract_text(ext: str, content: bytes) -> str:
    if ext in {".txt", ".md"}:
        return content.decode("utf-8", errors="ignore")
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    return ""


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ValueError("PDF support is not installed (pypdf)") from exc

    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValueError("That PDF is password protected") from exc
        pages = list(reader.pages[:MAX_PDF_PAGES])
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("That PDF could not be opened") from exc

    layout_parts: list[str] = []
    plain_parts: list[str] = []
    for page in pages:
        # Layout mode keeps two-column resumes from interleaving, but it bails
        # on some generators, so keep the default extraction as a fallback.
        try:
            layout_parts.append(page.extract_text(extraction_mode="layout") or "")
        except Exception:
            layout_parts.append("")
        try:
            plain_parts.append(page.extract_text() or "")
        except Exception:
            plain_parts.append("")

    layout = "\n".join(layout_parts).strip()
    plain = "\n".join(plain_parts).strip()
    if len(layout) >= 0.8 * len(plain):
        return layout
    return plain


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValueError("DOCX support is not installed (python-docx)") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError("That DOCX could not be opened") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]

    # Plenty of resume templates lay everything out in a table, which the
    # paragraph list alone reports as empty.
    for table in document.tables:
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                value = cell.text.strip()
                # Merged cells repeat their text across the span.
                if value and (not cells or cells[-1] != value):
                    cells.append(value)
            if cells:
                parts.append(" | ".join(cells))

    for section in document.sections:
        for holder in (section.header, section.footer):
            for paragraph in holder.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)

    return "\n".join(parts)


_LIGATURES = {"ﬁ": "fi", "ﬂ": "fl", "ﬀ": "ff", "ﬃ": "ffi", "ﬄ": "ffl"}


def _normalize_text(raw: str) -> str:
    if not raw:
        return ""
    text = unicodedata.normalize("NFKC", raw)
    for bad, good in _LIGATURES.items():
        text = text.replace(bad, good)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ").replace("\t", " ")
    text = re.sub(r"[\u2010-\u2015\u2212]", "-", text)
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    # Bullet glyphs start a new item rather than joining the previous word.
    text = re.sub(r"[\u2022\u25aa\u25cf\u25e6\u2023\u2043\u00b7\u25a0\u2219]", "\n", text)
    # Undo hyphenation across a line break: "Kuber-\nnetes".
    text = re.sub(r"(\w)-\n\s*(\w)", r"\1\2", text)
    # A wide run of spaces is a column gap in layout-extracted PDFs; treating it
    # as a line break keeps section headers on their own line.
    text = re.sub(r" {3,}", "\n", text)
    text = re.sub(r" {2}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.rstrip() for line in text.split("\n")).strip()


# --------------------------------------------------------------------------
# Sectioning
# --------------------------------------------------------------------------


def _match_section_keyword(core: str) -> str | None:
    cleaned = core.strip().strip(":").strip()
    if not cleaned:
        return None
    for name, pattern in SECTION_PATTERNS:
        if re.fullmatch(pattern, cleaned, flags=re.I):
            return name
    return None


def _classify_line(line: str) -> tuple[str | None, str]:
    """Return (section name, text that followed on the same line)."""
    raw = line.strip()
    if not raw or len(raw) > 120:
        return None, ""

    # "SKILLS: Python, Go" - header and content share a line.
    inline = re.match(r"^([A-Za-z][A-Za-z &/'()+.-]{1,46})\s*:\s*(\S.*)$", raw)
    if inline:
        name = _match_section_keyword(inline.group(1))
        return (name, inline.group(2).strip()) if name else (None, "")

    core = raw.strip(" :|*#-")
    if not core or len(core) > 50 or len(core.split()) > 5:
        return None, ""
    if not re.fullmatch(r"[A-Za-z][A-Za-z &/'()+,.-]*", core):
        return None, ""

    letters = [c for c in core if c.isalpha()]
    is_upper = bool(letters) and all(c.isupper() for c in letters)
    if not (is_upper or raw.endswith(":") or core.istitle()):
        return None, ""

    name = _match_section_keyword(core)
    return (name, "") if name else (None, "")


def _split_sections(text: str) -> dict[str, str]:
    buckets: dict[str, list[str]] = {"_header": []}
    current = "_header"
    for line in text.split("\n"):
        name, inline = _classify_line(line)
        if name:
            current = name
            buckets.setdefault(current, [])
            if inline:
                buckets[current].append(inline)
            continue
        buckets.setdefault(current, []).append(line)
    return {name: "\n".join(lines).strip() for name, lines in buckets.items()}


def _contact_block(text: str, sections: dict[str, str]) -> str:
    """The name / email / phone / city lines at the very top."""
    source = sections.get("_header") or text
    lines = [line for line in source.split("\n") if line.strip()]
    return "\n".join(lines[:6])


# --------------------------------------------------------------------------
# Roles
# --------------------------------------------------------------------------


def _canonical_seniority(raw: str | None) -> str:
    if not raw:
        return ""
    return SENIORITY_FORMS.get(raw.strip().lower().rstrip("."), "")


def _line_containing(text: str, index: int) -> str:
    start = text.rfind("\n", 0, index) + 1
    end = text.find("\n", index)
    return text[start:] if end == -1 else text[start:end]


def _extract_roles(text: str, sections: dict[str, str]) -> list[str]:
    scores: dict[str, float] = {}
    seniority_score: dict[str, dict[str, float]] = {}

    for name, body in sections.items():
        if not body:
            continue
        weight = ROLE_SECTION_WEIGHT.get(name, 1.0)
        for match in TITLE_RE.finditer(body):
            canonical = TITLE_CANONICAL[match.group("title").lower()]
            line = _line_containing(body, match.start())
            hit = weight

            # A line that is mostly the title is a heading, not prose.
            stripped = line.strip()
            if stripped and len(stripped) <= 80:
                density = len(match.group(0)) / len(stripped)
                if density >= 0.35:
                    hit *= 1.6

            # "Worked with Product Manager ..." describes a colleague.
            if THIRD_PARTY_CONTEXT.search(line):
                hit *= 0.15

            scores[canonical] = scores.get(canonical, 0.0) + hit
            level = _canonical_seniority(match.group("sen"))
            if level:
                bucket = seniority_score.setdefault(canonical, {})
                bucket[level] = bucket.get(level, 0.0) + hit

    seeking = _seeking_title(sections)

    # One passing mention in a bullet is not a career. Keep only titles with
    # real backing, unless that would leave us with nothing at all.
    strong = {k: v for k, v in scores.items() if v >= MIN_ROLE_SCORE}
    if strong:
        scores = strong

    # Prefer the more specific title when both are present.
    for canonical in list(scores):
        scores[canonical] += 0.5 * len(canonical.split())

    if seeking:
        scores[seeking] = scores.get(seeking, 0.0) + 25.0

    if not scores:
        return []

    specific = {k: v for k, v in scores.items() if k not in GENERIC_TITLES}
    if specific:
        scores = specific

    ranked = sorted(scores, key=lambda k: (-scores[k], k))
    primary = ranked[0]

    level = ""
    if primary in seniority_score and primary != seeking:
        level = max(seniority_score[primary].items(), key=lambda kv: kv[1])[0]

    role = f"{level} {primary}".strip()
    alternate = ranked[1] if len(ranked) > 1 else (primary if level else "")
    return [role, alternate] if alternate else [role]


def _seeking_title(sections: dict[str, str]) -> str:
    """A title from an explicit 'Seeking: ...' line, if it really is a title."""
    for name in ("_header", "summary"):
        body = sections.get(name, "")
        if not body:
            continue
        for match in SEEKING_RE.finditer(body):
            candidate = _clean_seeking_candidate(match.group(1))
            if candidate:
                return candidate
    return ""


def _clean_seeking_candidate(raw: str) -> str:
    candidate = raw.strip()
    candidate = re.split(r"[|•;(]|\s+(?:at|in|with|for|to)\s+", candidate, maxsplit=1)[0]
    candidate = re.split(r"\s*(?:,|/|\bor\b|\band\b)\s*", candidate, maxsplit=1, flags=re.I)[0]
    candidate = candidate.strip(" .,:;-|/")
    candidate = re.sub(
        r"\s+(?:roles?|positions?|jobs?|opportunit\w*|openings?)$", "", candidate, flags=re.I
    )
    if not candidate or len(candidate) > 45:
        return ""

    words = candidate.split()
    if not 1 <= len(words) <= 5:
        return ""
    # "open to remote opportunities" is a preference, not a job title.
    if not any(word.strip(".").lower() in ROLE_NOUNS for word in words[-2:]):
        return ""

    match = TITLE_RE.search(candidate)
    if match and match.group(0).strip().lower() == candidate.lower():
        canonical = TITLE_CANONICAL[match.group("title").lower()]
        level = _canonical_seniority(match.group("sen"))
        return f"{level} {canonical}".strip()
    if match:
        return TITLE_CANONICAL[match.group("title").lower()]
    return " ".join(word.capitalize() if word.islower() else word for word in words)


# --------------------------------------------------------------------------
# Skills
# --------------------------------------------------------------------------


def _extract_skills(text: str, sections: dict[str, str]) -> list[str]:
    skills_section = sections.get("skills", "")
    scores: dict[str, float] = {}

    for canonical, matchers in SKILL_MATCHERS:
        score = 0.0
        for _form, pattern, ambiguous in matchers:
            if skills_section and pattern.search(skills_section):
                score += 4.0
            if ambiguous:
                # Only trust these inside an explicit skills list.
                continue
            hits = len(pattern.findall(text))
            if hits:
                score += 1.0 + min(hits - 1, 3) * 0.5
        if score:
            scores[canonical] = score

    for child, parent in SKILL_SUPERSEDES.items():
        if child in scores and parent in scores:
            scores.pop(parent, None)

    order = {canonical: i for i, (canonical, _) in enumerate(_SKILL_DEFS)}
    ranked = sorted(scores, key=lambda k: (-scores[k], order[k]))[:MAX_SKILLS]

    if len(ranked) < MAX_SKILLS and skills_section:
        ranked.extend(_unlisted_skills(skills_section, ranked, MAX_SKILLS - len(ranked)))
    return ranked


_EXTRA_SKILL_STOPWORDS = re.compile(
    r"\b(?:and|or|with|strong|excellent|good|basic|advanced|intermediate|"
    r"proficien\w*|knowledge|experience|years?|ability|able|skills?|etc|"
    r"including|familiar\w*|working|hands|team|communication|leadership|"
    r"problem|solving|management|written|verbal|fluent|native|understanding|"
    r"expertise|exposure|concepts?|fundamentals?)\b",
    re.I,
)


def _unlisted_skills(section: str, already: list[str], limit: int) -> list[str]:
    """Pick up stack items from the skills section that aren't in our list."""
    taken = {_skill_key(name) for name in already}
    extras: list[str] = []
    for chunk in re.split(r"[,;|\n/]+", section):
        # Leading "." is part of the name (".NET"), so only trim the right side.
        item = chunk.lstrip(" :([").rstrip(" .,:;-)]")
        if not item or len(item) > 28 or len(item) < 2:
            continue
        if len(item.split()) > 3 or not re.search(r"[A-Za-z]", item):
            continue
        if _EXTRA_SKILL_STOPWORDS.search(item):
            continue
        # Real tooling is capitalised or carries a symbol; prose is not.
        if not (any(c.isupper() for c in item) or re.search(r"[.+#/\-0-9]", item)):
            continue
        key = _skill_key(item)
        if not key or key in taken or key in ALL_SKILL_KEYS:
            continue
        taken.add(key)
        extras.append(item)
        if len(extras) >= limit:
            break
    return extras


# --------------------------------------------------------------------------
# Years of experience
# --------------------------------------------------------------------------


def _month_index(year: int, month: int | None) -> int:
    # Bare years are assumed to be mid-year, so "2018 - 2024" is 6 years rather
    # than the 7 you would get from January-to-December.
    return year * 12 + (month if month else 6)


def _month_number(name: str | None, numeric: str | None) -> int | None:
    if name:
        return _MONTH_NUMS.get(name.strip(". ").lower())
    if numeric:
        value = int(numeric)
        if 1 <= value <= 12:
            return value
    return None


def _now_index() -> int:
    today = date.today()
    return today.year * 12 + today.month


def _date_ranges(text: str) -> list[tuple[int, int]]:
    now = _now_index()
    ranges: list[tuple[int, int]] = []
    for match in DATE_RANGE_RE.finditer(text):
        year1 = int(match.group("y1"))
        if not 1980 <= year1 <= 2035:
            continue
        start = _month_index(year1, _month_number(match.group("m1"), match.group("n1")))
        if match.group("present"):
            end = now
        else:
            year2 = int(match.group("y2"))
            if not 1980 <= year2 <= 2035:
                continue
            end = _month_index(year2, _month_number(match.group("m2"), match.group("n2")))
        end = min(end, now)
        if end <= start:
            continue
        ranges.append((start, end))
    return ranges


def _merge_months(ranges: list[tuple[int, int]]) -> int:
    """Total months covered, counting overlapping jobs only once."""
    if not ranges:
        return 0
    ordered = sorted(ranges)
    total = 0
    current_start, current_end = ordered[0]
    for start, end in ordered[1:]:
        if start <= current_end:
            current_end = max(current_end, end)
        else:
            total += current_end - current_start
            current_start, current_end = start, end
    return total + (current_end - current_start)


def _employment_text(text: str, sections: dict[str, str]) -> str:
    experience = sections.get("experience", "")
    if experience:
        return experience
    # No recognisable experience section: use everything except the parts whose
    # dates are not employment (degrees, certificates, awards).
    skip = {"education", "certifications", "other", "projects"}
    body = "\n".join(
        value for name, value in sections.items() if name not in skip and value
    )
    return "\n".join(
        line for line in (body or text).split("\n") if not EDUCATION_LINE.search(line)
    )


def _stated_years(sections: dict[str, str]) -> float | None:
    """Years the candidate claims outright, read only from their own summary."""
    blocks = [sections.get("_header", ""), sections.get("summary", "")]
    best: float | None = None
    for block in blocks:
        if not block:
            continue
        for line in block.split("\n"):
            if THIRD_PARTY_CONTEXT.search(line):
                continue
            for regex in (EXPLICIT_YEARS_RE, LABELLED_YEARS_RE):
                for match in regex.finditer(line):
                    value = float(match.group(1))
                    if 0 <= value <= 45 and (best is None or value > best):
                        best = value
    return best


def _extract_years(text: str, sections: dict[str, str]) -> float | None:
    computed: float | None = None
    months = _merge_months(_date_ranges(_employment_text(text, sections)))
    if months > 0:
        computed = round(months / 12, 1)

    stated = _stated_years(sections)
    if stated is None:
        return computed
    if computed is None:
        return stated
    # A written claim usually beats date arithmetic, but resumes go stale, so
    # let a clearly longer history win.
    return computed if computed > stated + 1.5 else stated


def _years_to_bands(years: float) -> list[str]:
    for band_id, low, high in EXPERIENCE_BANDS:
        if low <= years < high or (band_id == "12+" and years >= low):
            return [band_id]
    return []


# --------------------------------------------------------------------------
# Locations
# --------------------------------------------------------------------------


def _scan_locations(block: str, found: list[dict], seen: set[str]) -> None:
    spans: list[tuple[int, int]] = []
    ordered = sorted(KNOWN_LOCATIONS, key=lambda item: len(item[0]), reverse=True)
    for label, meta in ordered:
        for match in re.finditer(rf"(?<![A-Za-z]){re.escape(label)}(?![A-Za-z])", block, re.I):
            # "New Delhi" already claimed these characters; skip the "Delhi" hit.
            if any(start <= match.start() < end for start, end in spans):
                continue
            if INSTITUTION_SUFFIX.match(block[match.end():match.end() + 40]):
                continue
            key = meta.get("city") or meta.get("country") or label
            if key in seen:
                continue
            spans.append((match.start(), match.end()))
            seen.add(key)
            entry = {"label": label, **meta}
            entry.setdefault("remote", False)
            found.append(entry)
            break
        if len(found) >= 3:
            return


def _extract_locations(text: str, sections: dict[str, str]) -> list[dict]:
    found: list[dict] = []
    seen: set[str] = set()

    _scan_locations(_contact_block(text, sections), found, seen)

    if not found:
        for match in LOCATION_LABEL_RE.finditer(text):
            _scan_locations(match.group(1), found, seen)
            if found:
                break

    if not found:
        _scan_locations(text[:1500], found, seen)

    # A city already implies its country, so drop the bare country entry.
    cities = {entry["country"] for entry in found if entry.get("city")}
    return [
        entry
        for entry in found
        if entry.get("city") or entry.get("remote") or entry["country"] not in cities
    ]


# --------------------------------------------------------------------------
# Summary
# --------------------------------------------------------------------------


def _short_summary(text: str, profile: ResumeProfile) -> str:
    bits: list[str] = []
    if profile.role:
        bits.append(profile.role)
    if profile.experience_years is not None:
        bits.append(f"{profile.experience_years:g} yrs")
    if profile.skills:
        bits.append(", ".join(profile.skills[:5]))
    if profile.locations:
        bits.append(profile.locations[0]["label"])
    if bits:
        return " · ".join(bits)
    stripped = text.strip()
    return stripped.splitlines()[0][:120] if stripped else ""
